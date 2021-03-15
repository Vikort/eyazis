import time
import nltk
from nltk.corpus import wordnet as wn
from nltk.draw import TreeWidget
from nltk.draw.util import CanvasFrame
from tkinter import messagebox, Label, RIGHT, Frame, Tk, LEFT, Text, WORD, END, Menu, Button
from tkinter.filedialog import asksaveasfilename, askopenfilename
from docx import Document

nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('universal_tagset')

root = Tk()
root.option_add('*Dialog.msg.font', 'Helvetica 10')
root.title("Syntax parse tree")
root.resizable(width=False, height=False)

label_frame = Frame(root)
label_frame.pack()

enter_label = Label(label_frame, text='Введите текст', width=10)
enter_label.pack(side=LEFT)

empty_label = Label(label_frame, text='', width=65)
empty_label.pack(side=RIGHT)

enter_text = Text(width=70, height=20, wrap=WORD)
enter_text.pack()

canvas = CanvasFrame(root, width=550, height=180)
canvas.pack()


def docx_parser(docx):
    doc = Document(docx)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text


def open_file_and_input_text():
    file_name = askopenfilename(filetypes=(("Docx files", "*.docx"),))
    if file_name != '':
        text = docx_parser(file_name)
        enter_text.delete(1.0, END)
        enter_text.insert(1.0, text)


def save_docx():
    file = asksaveasfilename(filetypes=(("Docx file", "*.docx"),), defaultextension=("Docx file", "*.docx"))
    doc = Document()

    doc.add_paragraph(enter_text.get(1.0, END))
    doc.save(file)


help_info = """Система семантического анализа естественного языка

Система позволяет провести семантический анализ предложения анлглийского языка, загрузить предложение из файла, а также сохранить предложение в файл в формате docx.

Результат семантического анализ предложения представлен в виде дерева, узлами которого являются лексемы и их определение, синонимы, антонимы, гипонимы и гиперонимы.

Для проведения семантического анализа необходимо ввести текст в верхнее поле и затем нажать кнопку "Создать".
Для сохранение необходимо нажать кнопку "Сохранить", в появившемся окне выбрать нужный файл или задать имя новому файлу.
Для открытия словаря необходимо нажать пункт меню "Файл", в появившемся окне выбрать нужный файл.
"""


def information():
    messagebox.askquestion("Help", help_info, type='ok')


grammar = r"""
        P: {<PRT|ADP>}
        V: {<VERB>}
        N: {<NOUN|PRON>}
        NP: {<N|NP|P>+<ADJ|NUM|DET>+}
        NP: {<ADJ|NUM|DET>+<N|NP|P>+}
        PP: {<P><NP>|<NP><P>}
        VP: {<NP|N><V>}
        VP: {<VP><NP|N||ADV>}
        VP: {<NP|N|ADV><VP>}
        VP: {<VP><PP|P>}
        """


def draw_semantic_tree():
    canvas.canvas().delete("all")
    start = time.time()
    text = enter_text.get(1.0, END)
    text = text.replace('\n', '')
    text = text.replace(',', '')
    text = text.replace('.', '')
    if text != '':
        doc = nltk.word_tokenize(text)
        result = '(S '
        for word in doc:
            result += get_word_semantic(word)
        result += ')'
        result = nltk.tree.Tree.fromstring(result)
        widget = TreeWidget(canvas.canvas(), result)
        canvas.add_widget(widget, 50, 10)

    finish = time.time()
    delta = finish - start
    print('draw tree: ', delta)


def get_word_semantic(word: str) -> str:
    start = time.time()
    result = '(WS (W ' + word + ') (DEF ' + wn.synsets(word)[0].definition().replace(' ', '_') + ')'
    synonyms = []
    antonyms = []
    hyponyms = []
    hypernyms = []
    word = wn.synsets(word)
    for synset in word:
        for lemma in synset.lemmas():
            synonyms.append(lemma.name())
            if lemma.antonyms():
                antonyms.append(lemma.antonyms()[0].name())
    for hyponym in word[0].hyponyms():
        hyponyms.append(hyponym.name())
    for hypernym in word[0].hypernyms():
        hypernyms.append(hypernym.name())
    if len(synonyms):
        result += ' (SYN '
        for synonym in synonyms:
            result += synonym + ' '
    if len(antonyms):
        result += ') (ANT '
        for antonym in antonyms:
            result += antonym + ' '
    if len(hyponyms):
        result += ') (HY '
        for hyponym in hyponyms:
            result += hyponym + ' '
    if len(hypernyms):
        result += ') (HE '
        for hypernym in hypernyms:
            result += hypernym + ' '
    result += '))'
    print('get word semantic', time.time() - start)
    return result


main_menu = Menu(root)
main_menu.add_command(label='Файл', command=open_file_and_input_text)
main_menu.add_command(label='Помощь', command=information)
root.config(menu=main_menu)

button1 = Button(text="Создать", command=draw_semantic_tree)
button1.pack(side=LEFT)

button2 = Button(text="Сохранить", command=save_docx)
button2.pack(side=RIGHT)

root.mainloop()
